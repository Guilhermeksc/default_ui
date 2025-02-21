# imports
import re
import math
import pandas as pd
import csv

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from datetime import datetime, timedelta
import numpy as np
import os
from babel.dates import format_date
import locale
#from config import *

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_NUMERIC, 'pt_BR.UTF-8')


def resolve_data(data: str) -> str:
    """Função para parsear timestamp em string para data formatada em string.
    Exemplo:
    data = "2020-07-31 12:23:15"
    resove_data(data) == '31/07/2020'
    """
    data_tempo = datetime.strptime(data, "%Y-%m-%d %H:%M:%S")
    #data_tempo = datetime.strptime(data, "%m%Y")
    return data_tempo.strftime("%d/%m/%Y")

def replace_string(modelo: Document, flag: str, novo_texto: str) -> Document:
    """Função para encontrar uma flag no texto e substituir pelo valor de uma variável"""
    if isinstance(novo_texto, str):
        new_text = novo_texto
    else:
        new_text = novo_texto[0]

    doc = modelo
    for paragrafo in doc.paragraphs:
        if flag in paragrafo.text:
            inline = paragrafo.runs
            for i in range(len(inline)):
                texto = inline[i].text
                if flag in texto:
                    texto_substituido = texto.replace(flag, new_text)
                    inline[i].text = texto_substituido

    section = modelo.sections[0]
    header = section.header

    for paragrafo in header.paragraphs:
        if flag in paragrafo.text:
            inline = paragrafo.runs
            for i in range(len(inline)):
                texto = inline[i].text
                if flag in texto:
                    texto_substituido = texto.replace(flag, new_text)
                    inline[i].text = texto_substituido

    return modelo

def move_table_after(table, paragraph):
    """Move uma tabela para imediatamente abaixo do parágrafo indicado"""
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def insert_nip(modelo: Document, dados: pd.Series, titulo: str) -> Document:
    """Transforma uma Série do pandas em uma tabela com 6 colunas e insere no documento abaixo da flag"""
    flag = "<tabela_abaixo>"
    paragrafos = modelo.paragraphs
    linhas_tabela = [dados[titulo][i:i+6] for i in range(0, len(dados), 6)]

    # Perrcorre os parágrafos do documento, encontra
    #a flag para tabela, insere a tabela no local e apaga a flag
    for paragrafo in paragrafos:
        if flag in paragrafo.text:
            index = paragrafos.index(paragrafo)
            replace_string(modelo, flag, "")
            break
    local_tabela = paragrafos[index]

    # Adiciona uma tabela com 6 colunas e comprimento igual
    #ao número de linhas desejado
    table = modelo.add_table(cols=6, rows=len(linhas_tabela) +1)
    table.style = 'Table Grid'

    # Mescla as células da primeira linha
    for label_cell in range(1, 6):
        table.cell(0, 0).merge(table.cell(0, label_cell))

    # Insere o Título da Tabela
    titulo_tabela = table.cell(0, 0).paragraphs[0]
    titulo_tabela.text = titulo
    titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_tabela.style = modelo.styles['Título da tabela']

    # Percorre as células da tabela inserindo o nip correspondente
    for row_index in range(len(linhas_tabela)):
        for cell, nip in zip(table.row_cells(row_index+1), linhas_tabela[row_index]):
            paragrafo = cell.paragraphs[0]
            paragrafo.text = str(nip)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Inserir tracinhos nas células vazias
    for cell in table.rows[-1].cells:
        if cell.paragraphs[0].text == '':
            cell.paragraphs[0].text = '- - - - -'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    move_table_after(table, local_tabela)
    return modelo


def insert_table(modelo: Document, dados: pd.DataFrame) -> Document:
    """Transforma um dataframe do pandas em uma tabela e insere no documento abaixo da flag"""
    flag = "<tabela_abaixo>"

    valor_total = dados['Total2'].astype(float).sum()
    valor_total = locale.currency(float(valor_total), grouping=True)
    valor_total = 'TOTAL ' + valor_total

    dados = dados.drop(columns=['Total2'])

    linhas = dados.shape[0] + 2
    colunas = dados.shape[1]

    paragrafos = modelo.paragraphs
    for paragrafo in paragrafos:
        if flag in paragrafo.text:
            index = paragrafos.index(paragrafo)
            replace_string(modelo, flag, "")
            break

    local_tabela = paragrafos[index]
    tabela = modelo.add_table(linhas, colunas)
    tabela.style = 'Table Grid'

    # Inclui os títulos das colunas
    for label_index in range(dados.shape[1]):
        paragrafo = tabela.cell(0, label_index).paragraphs[0]
        paragrafo.text = dados.columns[label_index]
        paragrafo.style = modelo.styles['Título da tabela']
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adiciona os dados do dataframe nas células
    #correspondentes da tabela criada
    for linha in range(dados.shape[0]):
        for coluna in range(dados.shape[-1]):
            paragrafo = tabela.cell(linha+1, coluna).paragraphs[0]
            paragrafo.text = str(dados.values[linha, coluna])
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragrafo = tabela.cell(linhas - 1, 0).paragraphs[0]
    paragrafo.text = valor_total
    paragrafo.style = modelo.styles['Título da tabela']
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # insere célula do valor total
    for label_cell in range(1, colunas):
        tabela.cell(linhas - 1, 0).merge(tabela.cell(linhas - 1, label_cell))

    move_table_after(tabela, local_tabela)
    return modelo

def main(df_inconsistencias, path_notas='./notas/',
         coluna_oms=None, path_modelo=None,
         coluna_nips=None, titulo=None,
         num_nota=None, nome_notas=None ):

    df = df_inconsistencias

    for om in df[coluna_oms].unique():
        assunto = 'munic'
        dados_om = df[df[coluna_oms] == om]
        sigla = list(dados_om['OC'])[0]
        nome_nota = nome_notas.format( sigla=sigla)
        modelo = Document(os.path.join(path_modelo))


        dados = dados_om.drop(columns=[coluna_oms])

        replace = {'<sigla_om>': str(sigla),
                   # '<valor>': valor_total,
                    #'<UG>': ug,
                    #'<num_nota>': str(seq),
                    '<data_hoje>': data_hoje,
                    '<mmaaa>': data_MMAAAA,
                    '<_prazo>': prazo
                    }

        if titulo is None:
            modelo = insert_table(modelo=modelo, dados=dados)
        else:
            modelo = insert_nip(modelo=modelo, dados=dados,
                                titulo=titulo)

        for old, novo_texto in replace.items():
            modelo = replace_string(modelo=modelo, flag=old,
                                    novo_texto=novo_texto)
        modelo.save(path_notas + f'{nome_nota}.docx')



data = datetime.now()
data_hoje = format_date(data, 'long', locale='pt_BR')
ficha = data - timedelta(30)
data_MMAAAA = ficha.strftime('%B%Y').capitalize()
prazo = (data + timedelta(days=60)).strftime('%d%b%Y').upper()

# Argumentos que sempre mudam
nome_planilha = './MUNIC_EM_LIC.xlsx'
nome_aba_base = 'Dados NA'
#nome_aba_siglas = 'Siglas'
colunas_valor = ['VALOR TOTAL']
colunas_data = ['DATA INÍCIO' ,'DATA FIM','DATA INÍCIO DA LICENÇA']

args = {
    'titulo': None,
    'nome_notas': 'NA-{sigla}-munic_em_lic',
    'coluna_oms': "OC",
    'path_modelo': 'modelo-formatado.docx',
    'num_nota': 'S/N',
}


####### Daqui para baixo ########
# inserir nomes do arquivo e da aba com os dados

df_inconsistencias = pd.read_excel(nome_planilha,sheet_name=nome_aba_base, dtype=str)
df_inconsistencias.fillna('', inplace = True)

#if nome_aba_siglas:
#    df_siglas = pd.read_excel(nome_planilha,
#                                  sheet_name=nome_aba_siglas, dtype=str)
#    df_inconsistencias = pd.merge(df_inconsistencias, df_siglas, on=['uge'])

args['df_inconsistencias']= df_inconsistencias

####### Transformação de colunas
# Coluna com valor monetário (não esquecer do ponto no excel!!!)
# Inserir nomes das colunas com valor monetário na lista
for col in colunas_valor:
     df_inconsistencias[col] = df_inconsistencias[col].apply(lambda x: locale.currency(float(x), grouping=True))



# Se as colunas com data aparecerem mal formatadas no documento final:
# Inserir nomes das colunas com data na lista
for coluna in colunas_data:
    df_inconsistencias[coluna] = df_inconsistencias[coluna].apply(resolve_data)




main(**args)




# In[ ]:




